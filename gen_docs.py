from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return table

# ===== TITLE PAGE =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TÀI LIỆU DỰ ÁN\nPBMed - HỆ THỐNG QUẢN LÝ BỆNH VIỆN')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Hoàng Minh Nhất — MSSV 20237468\nĐồ án I — Hệ thống thông tin y tế\nĐại học Bách khoa Hà Nội')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run('Tài liệu được tạo ngày 23/06/2026')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ===== TABLE OF CONTENTS =====
doc.add_heading('MỤC LỤC', level=1)
toc_items = [
    '1. Tổng quan dự án',
    '2. Kiến trúc & Luồng dữ liệu',
    '3. Cấu trúc thư mục & file',
    '4. Core Logic — Phân tích chuyên sâu',
    '   4.1. Authentication Module (users/views.py)',
    '   4.2. Doctor Module (doctors/views.py)',
    '   4.3. Receptionist Module (receptionist/views.py)',
    '   4.4. Patient Module (patients/views.py)',
    '   4.5. Admin Module (administration/views.py)',
    '   4.6. Seed Script (users/management/commands/setup.py)',
    '5. Data Models & Schemas',
    '6. API Reference',
    '7. Cấu hình & Môi trường',
    '8. Hạn chế & Nợ kỹ thuật',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(11)

doc.add_page_break()

# ===== SECTION 1: OVERVIEW =====
doc.add_heading('1. Tổng quan dự án', level=1)

doc.add_heading('1.1. Giới thiệu', level=2)
doc.add_paragraph(
    'PBMed là hệ thống quản lý bệnh viện mã nguồn mở được xây dựng trên nền tảng Django, '
    'hướng đến các cơ sở y tế vừa và nhỏ tại Việt Nam. Hệ thống hỗ trợ bốn vai trò người dùng: '
    'Bệnh nhân, Bác sĩ, Lễ tân và Quản trị viên, mỗi vai trò có một bộ chức năng riêng biệt '
    'phục vụ quy trình khám chữa bệnh cơ bản.'
)
doc.add_paragraph(
    'Hệ thống giải quyết bài toán quản lý tập trung các hoạt động: đặt lịch hẹn khám bệnh trực tuyến, '
    'tiếp nhận bệnh nhân tại quầy, tạo và quản lý phiếu khám bệnh, chỉ định và trả kết quả cận lâm sàng (CLS), '
    'và báo cáo thống kê cho quản trị. Dữ liệu được lưu trữ tập trung trên MySQL, đảm bảo tính nhất quán '
    'và an toàn thông tin y tế.'
)

doc.add_heading('1.2. Công nghệ sử dụng', level=2)
tech_headers = ['Cấu phần', 'Công nghệ', 'Lý do chọn']
tech_rows = [
    ['Backend Framework', 'Django 5.2', 'Kiến trúc MTV (Model-Template-View) mặc định, ORM mạnh, admin panel built-in'],
    ['Ngôn ngữ', 'Python 3.11+', 'Hệ sinh thái phong phú, dễ bảo trì, phù hợp ứng dụng web doanh nghiệp'],
    ['Cơ sở dữ liệu', 'MySQL 8/9', 'Ổn định, hỗ trợ UTF8MB4, phù hợp lưu trữ dữ liệu y tế có cấu trúc'],
    ['Template Engine', 'Django Templates', 'Tích hợp sẵn, không cần thư viện JS phức tạp'],
    ['UI Framework', 'Bootstrap 5', 'Responsive, dễ tùy biến, hỗ trợ tiếng Việt tốt'],
    ['Icons', 'Bootstrap Icons', 'Nhẹ, đồng bộ với Bootstrap'],
    ['Xác thực', 'Django Auth + Session', 'Mặc định, an toàn, tích hợp CSRF, session-based'],
]
add_table(doc, tech_headers, tech_rows)

doc.add_page_break()

# ===== SECTION 2: ARCHITECTURE =====
doc.add_heading('2. Kiến trúc & Luồng dữ liệu', level=1)

doc.add_heading('2.1. Kiến trúc tổng thể', level=2)
doc.add_paragraph(
    'Hệ thống áp dụng mô hình kiến trúc 3 tầng (Three-Tier Architecture) kết hợp với mẫu MTV '
    '(Model-Template-View) của Django:'
)
doc.add_paragraph('• Tầng Presentation (Giao diện): Templates HTML + Bootstrap 5, chịu trách nhiệm hiển thị dữ liệu và tương tác người dùng.', style='List Bullet')
doc.add_paragraph('• Tầng Business Logic (Xử lý nghiệp vụ): Django Views, chứa toàn bộ logic xử lý request, phân quyền, validation.', style='List Bullet')
doc.add_paragraph('• Tầng Data (Dữ liệu): Django Models ánh xạ tới các bảng MySQL qua ORM.', style='List Bullet')

doc.add_paragraph(
    'Các module được tổ chức theo tính năng (feature-based), mỗi app Django chịu trách nhiệm '
    'một nhóm chức năng cụ thể. Giao tiếp giữa các module thông qua Django ORM (đọc/ghi dữ liệu) '
    'và URL routing (điều hướng người dùng).'
)

doc.add_heading('2.2. Luồng dữ liệu điển hình — Đặt lịch hẹn', level=2)
doc.add_paragraph(
    '1. Bệnh nhân đăng nhập → gửi request GET tới /book_appointment/\n'
    '2. View book_appointment() query danh sách bác sĩ (Doctors) kèm chuyên khoa (Specialty) và khung giờ (Time)\n'
    '3. Template book_appointment.html render form với dữ liệu bác sĩ, chuyên khoa, lịch\n'
    '4. Bệnh nhân chọn bác sĩ, ngày, giờ → gửi POST tới /confirm_book/<doctor>/\n'
    '5. View patient_confirm_book() tạo Appointment mới với status="Waited"\n'
    '6. Appointments được lưu vào DB → Lễ tân có thể xem và xác nhận từ receptionist dashboard\n'
    '7. Sau khi xác nhận, status chuyển thành "Accepted" → Bác sĩ thấy trong danh sách chờ khám\n'
    '8. Bác sĩ tạo phiếu khám (ExaminationRecord) → status chuyển thành "Da kham"'
)

doc.add_heading('2.3. Quyết định thiết kế', level=2)
decisions = [
    ('Tại sao dùng MySQL thay vì SQLite?', 
     'Yêu cầu lưu trữ dữ liệu y tế có cấu trúc, hỗ trợ concurrent access, UTF8MB4 cho tiếng Việt. '
     'SQLite không phù hợp cho production với nhiều người dùng đồng thời.'),
    ('Tại sao dùng Session Auth thay vì JWT?',
     'Hệ thống là web app truyền thống (không phải SPA/API). Django session tích hợp sẵn, an toàn hơn '
     'với CSRF protection, dễ triển khai mà không cần frontend framework.'),
    ('Tại sao tách riêng PatientProfile khỏi User model?',
     'Một người dùng (User) có thể có nhiều hồ sơ bệnh nhân (PatientProfile) — ví dụ: đặt lịch cho '
     'người thân. Điều này linh hoạt hơn so với gộp chung vào User.'),
    ('Tại sao dùng decorator @doctor_required thay vì permission classes?',
     'Django REST Framework không được sử dụng. Decorator là cách đơn giản nhất để kiểm tra quyền '
     'trong view function-based, dễ đọc và dễ maintain.'),
    ('Tại sao seed data bằng Django ORM thay vì SQL?',
     'SQL thuần dễ bị lệch schema khi model thay đổi. ORM tự động khớp với model hiện tại, '
     'idempotent (chạy nhiều lần không duplicate), hỗ trợ tiếng Việt UTF-8 tốt hơn.'),
]
for title, desc in decisions:
    doc.add_paragraph(title, style='List Bullet')
    doc.add_paragraph(f'   → {desc}')

doc.add_page_break()

# ===== SECTION 3: FILE STRUCTURE =====
doc.add_heading('3. Cấu trúc thư mục & file', level=1)
doc.add_paragraph('Bảng dưới đây mô tả tất cả các file và thư mục có chức năng hoạt động trong hệ thống.')

file_headers = ['Đường dẫn', 'Mục đích']
file_rows = [
    ['manage.py', 'Entry point Django: chạy server, migrate, setup, shell'],
    ['.env', 'Biến môi trường: SECRET_KEY, DB_HOST, EMAIL config'],
    ['requirements.txt', 'Danh sách thư viện Python cần cài đặt'],
    ['.gitignore', 'Loại trừ file không cần commit (venv, .env, __pycache__)'],
    ['README.md', 'Hướng dẫn cài đặt và sử dụng dự án'],
    ['hospital/settings.py', 'Cấu hình Django: database, apps, middleware, template, static'],
    ['hospital/urls.py', 'Router URL chính: include 5 app + admin'],
    ['hospital/wsgi.py', 'WSGI entry point cho production server'],
    ['users/', 'App người dùng: xác thực, login/register, forgot/reset password'],
    ['users/models.py', 'Models: Users, Doctors, Patients, Specialty, Reste_token, ReceptionistProfile'],
    ['users/views.py', 'Views: register, login, forgot_view, reset_view, logout'],
    ['users/urls.py', 'URL routing cho authentication'],
    ['users/admin.py', 'Đăng ký models vào Django admin'],
    ['users/helpers.py', 'Helper: send_email() gửi email reset password'],
    ['users/management/commands/setup.py', 'Seed script: migrate + tạo dữ liệu mẫu bằng ORM'],
    ['patients/', 'App bệnh nhân: đặt lịch, hồ sơ, lịch sử khám'],
    ['patients/models.py', 'Models: Time, Status, PatientProfile, Appointment, CLSService, ExaminationRecord, CLSIndication'],
    ['patients/views.py', 'Views: book_appointment, my_appointments, profile CRUD, exam_history, cls_results'],
    ['patients/urls.py', 'URL routing cho bệnh nhân'],
    ['doctors/', 'App bác sĩ: dashboard, lịch hẹn, phiếu khám, CLS'],
    ['doctors/views.py', 'Views: doctor_dashboard, view_appointments, doctor_exam_create/edit, cls_pending_list, cls_update_result'],
    ['doctors/urls.py', 'URL routing cho bác sĩ'],
    ['receptionist/', 'App lễ tân: dashboard, walk-in, lịch hẹn'],
    ['receptionist/views.py', 'Views: dashboard, appointments, walkin, patient_list, schedule'],
    ['receptionist/urls.py', 'URL routing cho lễ tân'],
    ['administration/', 'App quản trị: dashboard, users, categories, reports'],
    ['administration/views.py', 'Views: dashboard, user CRUD, category CRUD, reports, CSV export'],
    ['administration/urls.py', 'URL routing cho quản trị'],
    ['templates/users/', 'Giao diện: login, register, forgot, reset password'],
    ['templates/patients/', 'Giao diện: book_appointment, my_appointments, profiles, exam_history'],
    ['templates/doctors/', 'Giao diện: dashboard, exam_form, exam_detail, cls_pending, patient_list'],
    ['templates/receptionist/', 'Giao diện: dashboard, appointments, walkin, patients, schedule'],
    ['templates/administration/', 'Giao diện: dashboard, users, categories, reports, category_edit, user_form'],
    ['static/', 'File tĩnh: CSS, JS, images, Bootstrap'],
]
add_table(doc, file_headers, file_rows)

doc.add_page_break()

# ===== SECTION 4: CORE LOGIC =====
doc.add_heading('4. Core Logic — Phân tích chuyên sâu', level=1)

# 4.1 Authentication
doc.add_heading('4.1. Authentication Module (users/views.py)', level=2)

doc.add_heading('register(request)', level=3)
doc.add_paragraph('Mục đích: Xử lý đăng ký tài khoản bệnh nhân mới. Cho phép người dùng tự tạo tài khoản với username, email, mật khẩu. Sau khi đăng ký thành công, tự động tạo bản ghi Patients liên kết.')
doc.add_paragraph('Parameters: request (HttpRequest) — chứa POST data: user_firstname, user_lastname, user_id (username), email, password, conf_password, user_gender, birthday.')
doc.add_paragraph('Returns: redirect đến login nếu thành công, render register form kèm error messages nếu thất bại.')
p = doc.add_paragraph('Edge cases:')
p.add_run('\n- Username đã tồn tại → báo lỗi "Tên đăng nhập đã tồn tại."')
p.add_run('\n- Email đã tồn tại → báo lỗi "Email đã được sử dụng."')
p.add_run('\n- Mật khẩu < 6 ký tự → báo lỗi + giữ lại form data')
p.add_run('\n- Mật khẩu không khớp → báo lỗi "Mật khẩu không khớp."')

doc.add_heading('login_view(request)', level=3)
doc.add_paragraph('Mục đích: Xác thực người dùng bằng username/password. Sau khi xác thực thành công, chuyển hướng đến dashboard tương ứng theo role.')
doc.add_paragraph('Flow: authenticate() → nếu user tồn tại → login(request, user) → redirect dựa trên is_doctor/is_admin/is_receptionist/Patients.exists()')

doc.add_heading('forgot_view(request) & reset_view(request, token)', level=3)
doc.add_paragraph('Mục đích: Xử lý quên mật khẩu. forgot_view nhận email, tạo token UUID, lưu Reste_token với expiry 24 giờ, gửi email (hoặc hiển thị link nếu email không gửi được). reset_view kiểm tra token, cho phép đặt mật khẩu mới.')
doc.add_paragraph('Security: Token được tạo bằng uuid.uuid4() (128-bit ngẫu nhiên). Token hết hạn sau 24h. Sau khi đặt lại mật khẩu, token bị xoá.')

# 4.2 Doctor Module
doc.add_heading('4.2. Doctor Module (doctors/views.py)', level=2)

doc.add_heading('doctor_required(view_func) — Decorator', level=3)
doc.add_paragraph('Mục đích: Decorator kiểm tra quyền bác sĩ. Chặn truy cập nếu user chưa đăng nhập hoặc không phải bác sĩ. Tự động redirect về trang login hoặc dashboard tương ứng.')
doc.add_paragraph('Giải pháp thay thế: Có thể dùng Group permissions của Django, nhưng decorator đơn giản hơn cho 4 vai trò riêng biệt (admin, doctor, receptionist, patient).')

doc.add_heading('doctor_dashboard(request)', level=3)
doc.add_paragraph('Mục đích: Trang dashboard của bác sĩ. Hiển thị thống kê (hôm nay/tuần/tháng) và luồng hoạt động gần đây (xác nhận lịch, hủy lịch, khám bệnh, kết quả CLS).')
doc.add_paragraph('Thống kê: examined_today, appointments_today, examined_week, examined_month, cls_month, total_examined, total_pending_cls — tất cả đều filter theo doctor hiện tại.')
doc.add_paragraph('Activity Feed: Gộp 4 query (confirm, cancel, exam, cls) vào list dict, sort theo thời gian, lấy 10 hoạt động gần nhất. Mỗi hoạt động có type, time, text, detail, color, icon.')

doc.add_heading('doctor_exam_create(request, appointment_id)', level=3)
doc.add_paragraph('Mục đích: Tạo phiếu khám mới cho một lịch hẹn đã được xác nhận (Accepted).')
doc.add_paragraph('Parameters: appointment_id (int) — ID của lịch hẹn.')
doc.add_paragraph('Flow:')
p = doc.add_paragraph()
p.add_run('1. Kiểm tra appointment tồn tại và thuộc về doctor hiện tại (get_object_or_404(..., doctor=doctor))\n')
p.add_run('2. Kiểm tra phiếu khám đã tồn tại → redirect đến detail nếu có\n')
p.add_run('3. Nếu POST: tạo ExaminationRecord với các trường từ form\n')
p.add_run('4. Xử lý CLS: nếu có cls_services từ form, tạo CLSIndication cho mỗi service\n')
p.add_run('5. Nếu appointment status != Accepted, đổi thành "Da kham" + set confirmed_at\n')
p.add_run('6. Redirect đến doctor_exam_detail')
doc.add_paragraph('Edge cases: appointment.patient_profile = None → báo lỗi "chưa có hồ sơ bệnh nhân".')

doc.add_heading('cls_pending_list(request) & cls_update_result(request, indication_id)', level=3)
doc.add_paragraph('Mục đích: Hiển thị danh sách chỉ định CLS đang chờ kết quả của bác sĩ hiện tại. Cho phép nhập kết quả xét nghiệm.')
doc.add_paragraph('Lưu ý bảo mật: Chỉ hiển thị CLS thuộc phiếu khám do bác sĩ hiện tại tạo — filter theo examination_record__doctor=doctor.')

# 4.3 Receptionist
doc.add_heading('4.3. Receptionist Module (receptionist/views.py)', level=2)

doc.add_heading('walkin(request)', level=3)
doc.add_paragraph('Mục đích: Đăng ký khám trực tiếp tại quầy cho bệnh nhân chưa có tài khoản. Tự động tạo tài khoản với username ngẫu nhiên (pbmed_<8hex>) — không dùng SĐT làm username để tránh lộ thông tin y tế khi SĐT bị thu hồi.')
doc.add_paragraph('Flow:')
p = doc.add_paragraph()
p.add_run('1. Kiểm tra patient profile theo SĐT → nếu có, dùng lại user cũ\n')
p.add_run('2. Nếu chưa có → tạo user mới với username random + password 123456 + Patients record\n')
p.add_run('3. Tạo PatientProfile (nếu chưa có profile khớp user+full_name)\n')
p.add_run('4. Tạo Appointment với status="Waited"')
doc.add_paragraph('Edge cases: phone đã tồn tại → dùng user cũ (không tạo mới). Phone chưa tồn tại → tạo mới. Cùng phone khác tên → tạo profile mới.')

doc.add_heading('appointments(request) & schedule(request)', level=3)
doc.add_paragraph('appointments: Danh sách lịch hẹn, hỗ trợ filter (status, date, doctor). Cho phép confirm (Accepted) và cancel (Cancelled).')
doc.add_paragraph('schedule: Xem lịch trực của bác sĩ trong tuần. Query thống kê số lịch hẹn theo từng bác sĩ.')

# 4.4 Patient Module
doc.add_heading('4.4. Patient Module (patients/views.py)', level=2)

doc.add_heading('book_appointment(request)', level=3)
doc.add_paragraph('Mục đích: Trang đặt lịch hẹn. Hiển thị danh sách bác sĩ kèm chuyên khoa, lịch 7 ngày tới, khung giờ sáng/chiều.')
doc.add_paragraph('Tách giờ: morning = [t for t in times if int(t.time) < 12], afternoon = [t for t in times if int(t.time) >= 12]')
doc.add_paragraph('Filter: theo chuyên khoa (specialty__name=fs) và tên bác sĩ (user__first_name__icontains=fn).')

doc.add_heading('patient_confirm_book(request, doctor)', level=3)
doc.add_paragraph('Mục đích: Xác nhận đặt lịch. Nhận profile_id, time, date, summary từ form → tạo Appointment với status="Waited".')
doc.add_paragraph('Parameters: doctor (str) — username của bác sĩ.')

doc.add_heading('exam_history & cls_results', level=3)
doc.add_paragraph('exam_history: Xem lịch sử khám bệnh theo profile. Query: ExaminationRecord.objects.filter(patient_profile=profile).')
doc.add_paragraph('cls_results: Xem kết quả CLS. Query: CLSIndication.objects.filter(examination_record__patient_profile=profile).')

# 4.5 Admin
doc.add_heading('4.5. Admin Module (administration/views.py)', level=2)

doc.add_heading('dashboard(request)', level=3)
doc.add_paragraph('Mục đích: Dashboard quản trị hiển thị tổng quan: số lượng doctor/patient/receptionist, appointments, phiếu khám, CLS — theo hôm nay, tháng, tổng.')
doc.add_paragraph('Top doctors: annotate total_exams=Count("exam_records"), order_by("-total_exams")[:5]')
doc.add_paragraph('Specialty stats: annotate doctor_count, exam_count → order_by("-exam_count").')

doc.add_heading('User CRUD (user_create, user_edit, user_delete)', level=3)
doc.add_paragraph('Mục đích: Quản lý tài khoản người dùng. Hỗ trợ tạo/sửa/xoá với 3 role: doctor, receptionist, patient.')
doc.add_paragraph('user_create: POST → kiểm tra username tồn tại → tạo User + tuỳ role tạo Doctors/Patients/ReceptionistProfile.')
doc.add_paragraph('user_edit: Sửa thông tin + mật khẩu + đổi role. Xử lý đổi role: xoá/cập nhật is_doctor/is_receptionist + tạo profile tương ứng.')
doc.add_paragraph('user_delete: Xoá user (CASCADE sẽ xoá các bản ghi liên quan).')

doc.add_heading('Category CRUD & Reports', level=3)
doc.add_paragraph('Category management: CRUD cho Specialty, CLSService, Status. Dùng _get_category_model() để map type string sang Model class.')
doc.add_paragraph('Reports: Thống kê 12 tháng (số appointments + exams theo tháng). Thống kê CLS service (total/completed/pending).')
doc.add_paragraph('CSV Export: 3 loại báo cáo (appointments, patients, examinations). Xuất file CSV có BOM cho Excel.')

# 4.6 Seed Script
doc.add_heading('4.6. Seed Script (setup.py)', level=2)
doc.add_paragraph('Mục đích: Script thiết lập database từ đầu — migrate + seed dữ liệu mẫu. Sử dụng Django ORM, idempotent (kiểm tra data tồn tại trước khi seed).')
doc.add_paragraph('Các bước:')
doc.add_paragraph('1. _seed_specialties(): Tạo 10 chuyên khoa tiếng Việt', style='List Bullet')
doc.add_paragraph('2. _seed_users(): Tạo admin + 5 bác sĩ + 2 lễ tân + 5 bệnh nhân', style='List Bullet')
doc.add_paragraph('3. _seed_profiles(): Tạo hồ sơ cho lễ tân', style='List Bullet')
doc.add_paragraph('4. _seed_catalogs(): Tạo trạng thái (Accepted/Waited/Cancelled/Da kham), khung giờ, dịch vụ CLS', style='List Bullet')
doc.add_paragraph('5. _seed_appointments(): Tạo 12 lịch hẹn + 6 phiếu khám + 12 chỉ định CLS + chi tiết cho doc001 (3 phiếu khám, 6 CLS)', style='List Bullet')
doc.add_paragraph('Lý do dùng ORM: Tránh lệch schema, xử lý FK constraint đúng thứ tự, hỗ trợ tiếng Việt, có thể chạy lại an toàn.')

doc.add_page_break()

# ===== SECTION 5: DATA MODELS =====
doc.add_heading('5. Data Models & Schemas', level=1)

doc.add_heading('5.1. users/models.py', level=2)

doc.add_heading('Users (AbstractUser)', level=3)
user_headers = ['Field', 'Type', 'Purpose', 'Constraints']
user_rows = [
    ['username', 'CharField(50)', 'Tên đăng nhập', 'unique=True'],
    ['email', 'CharField(50)', 'Email', 'unique=True'],
    ['password', 'CharField(200)', 'Mật khẩu (hashed)', ''],
    ['first_name', 'CharField(50)', 'Tên', ''],
    ['last_name', 'CharField(50)', 'Họ', ''],
    ['gender', 'CharField(10)', 'Giới tính', 'choices: Male/Female'],
    ['birthday', 'DateField', 'Ngày sinh', 'null=True'],
    ['is_doctor', 'BooleanField', 'Là bác sĩ?', 'default=False'],
    ['is_receptionist', 'BooleanField', 'Là lễ tân?', 'default=False'],
    ['is_admin', 'BooleanField', 'Là quản trị?', 'default=False'],
]
add_table(doc, user_headers, user_rows)
doc.add_paragraph('DB Table: system_user')

doc.add_heading('Doctors', level=3)
p = doc.add_paragraph('PK = user_id (OneToOne → Users). Fields: specialty (FK → Specialty), bio (Text). DB Table: doctor_profile.')

doc.add_heading('Patients', level=3)
p = doc.add_paragraph('PK = user_id (OneToOne → Users). Fields: insurance (CharField). DB Table: patient_account.')

doc.add_heading('Specialty', level=3)
p = doc.add_paragraph('Fields: name (unique), description. DB Table: specialty.')

doc.add_heading('Reste_token', level=3)
rt_headers = ['Field', 'Type', 'Purpose']
rt_rows = [
    ['user', 'FK -> Users', 'Người dùng yêu cầu reset'],
    ['email', 'CharField(50)', 'Email gửi link reset'],
    ['token', 'CharField(50)', 'UUID token'],
    ['created_at', 'DateTimeField', 'Thời gian tạo'],
    ['expiry_date', 'DateTimeField', 'Hết hạn sau 24h'],
]
add_table(doc, rt_headers, rt_rows)
doc.add_paragraph('DB Table: password_reset_token')

doc.add_heading('ReceptionistProfile', level=3)
p = doc.add_paragraph('PK = user_id (OneToOne → Users). Fields: phone, address, employee_id, start_date, note. DB Table: receptionist_profile.')

doc.add_heading('5.2. patients/models.py', level=2)

doc.add_heading('Time', level=3)
p = doc.add_paragraph('Field: time (CharField 10). Lưu chuỗi giờ như "08:00". DB Table: time_slot.')

doc.add_heading('Status', level=3)
p = doc.add_paragraph('Field: status (CharField 20). Giá trị: "Accepted", "Waited", "Cancelled", "Da kham". DB Table: appointment_status.')

doc.add_heading('PatientProfile', level=3)
pp_headers = ['Field', 'Type', 'Purpose']
pp_rows = [
    ['user', 'FK -> Users', 'Người dùng sở hữu hồ sơ'],
    ['full_name', 'CharField(100)', 'Họ tên bệnh nhân'],
    ['birth_year', 'IntegerField', 'Năm sinh'],
    ['birth_date', 'DateField', 'Ngày sinh (null allowed)'],
    ['gender', 'CharField(10)', 'Giới tính'],
    ['phone', 'CharField(15)', 'Số điện thoại'],
    ['cccd', 'CharField(20)', 'Căn cước công dân'],
    ['address', 'TextField', 'Địa chỉ'],
    ['ward', 'CharField(100)', 'Phường/Xã'],
    ['city', 'CharField(100)', 'Tỉnh/Thành phố'],
    ['insurance', 'CharField(50)', 'Mã BHYT'],
    ['created_at', 'DateTimeField', 'Ngày tạo (auto)'],
]
add_table(doc, pp_headers, pp_rows)
doc.add_paragraph('DB Table: patient_profile')

doc.add_heading('Appointment', level=3)
app_headers = ['Field', 'Type', 'Purpose']
app_rows = [
    ['doctor', 'FK -> Doctors', 'Bác sĩ phụ trách'],
    ['patient', 'FK -> Patients', 'Tài khoản bệnh nhân'],
    ['patient_profile', 'FK -> PatientProfile', 'Hồ sơ BN (null allowed)'],
    ['summary', 'TextField', 'Lý do khám'],
    ['description', 'TextField', 'Mô tả chi tiết'],
    ['start_date', 'DateField', 'Ngày hẹn'],
    ['time', 'FK -> Time', 'Khung giờ'],
    ['status', 'FK -> Status', 'Trạng thái'],
    ['rejection_reason', 'CharField(255)', 'Lý do hủy'],
    ['created_at', 'DateTimeField', 'Ngày tạo'],
    ['confirmed_at', 'DateTimeField', 'Ngày xác nhận'],
    ['rejected_at', 'DateTimeField', 'Ngày hủy'],
]
add_table(doc, app_headers, app_rows)
doc.add_paragraph('DB Table: appointment')
doc.add_paragraph('Status flow: Waited (chờ xác nhận) → Accepted (chờ khám) → Da kham (đã khám) | Cancelled (đã hủy)')

doc.add_heading('ExaminationRecord', level=3)
er_headers = ['Field', 'Type', 'Purpose']
er_rows = [
    ['patient_profile', 'FK -> PatientProfile', 'Bệnh nhân'],
    ['doctor', 'FK -> Doctors', 'Bác sĩ khám'],
    ['appointment', 'FK -> Appointment', 'Lịch hẹn liên kết'],
    ['exam_date', 'DateTimeField', 'Ngày giờ khám'],
    ['room', 'CharField(50)', 'Phòng khám'],
    ['reason', 'TextField', 'Lý do khám'],
    ['medical_history', 'TextField', 'Tiền sử bệnh'],
    ['family_history', 'TextField', 'Tiền sử gia đình'],
    ['pulse', 'CharField(20)', 'Mạch'],
    ['temperature', 'CharField(20)', 'Nhiệt độ'],
    ['blood_pressure', 'CharField(20)', 'Huyết áp'],
    ['weight', 'CharField(20)', 'Cân nặng'],
    ['health_status', 'TextField', 'Tình trạng sức khỏe'],
    ['diagnosis', 'TextField', 'Chẩn đoán'],
]
add_table(doc, er_headers, er_rows)
doc.add_paragraph('DB Table: examination_record')

doc.add_heading('CLSIndication', level=3)
cls_headers = ['Field', 'Type', 'Purpose']
cls_rows = [
    ['examination_record', 'FK -> ExaminationRecord', 'Phiếu khám'],
    ['cls_service', 'FK -> CLSService', 'Loại xét nghiệm'],
    ['status', 'CharField(20)', 'Trạng thái: pending/completed'],
    ['result', 'TextField', 'Kết quả'],
    ['result_date', 'DateTimeField', 'Ngày trả kết quả'],
    ['performed_by', 'FK -> Users', 'Người thực hiện'],
]
add_table(doc, cls_headers, cls_rows)
doc.add_paragraph('DB Table: cls_indication')

doc.add_heading('CLSService', level=3)
p = doc.add_paragraph('Fields: name (unique), description (Text). DB Table: cls_service.')
doc.add_paragraph('Giá trị mẫu: Xét nghiệm máu, Xét nghiệm nước tiểu, Siêu âm tổng quát, X-quang ngực, Điện tâm đồ, Sinh hóa máu, Nội soi tiêu hóa.')

doc.add_heading('5.3. Quan hệ giữa các Models', level=2)
rel_headers = ['Model 1', 'Quan hệ', 'Model 2', 'Giải thích']
rel_rows = [
    ['Users', '1-1', 'Doctors', 'Mỗi user là bác sĩ có 1 profile doctor'],
    ['Users', '1-1', 'Patients', 'Mỗi user là bệnh nhân có 1 tài khoản patient'],
    ['Users', '1-N', 'PatientProfile', 'Một user có thể tạo nhiều hồ sơ (cho người thân)'],
    ['Users', '1-1', 'ReceptionistProfile', 'Mỗi lễ tân có 1 hồ sơ'],
    ['Doctors', 'N-1', 'Specialty', 'Mỗi bác sĩ thuộc 1 chuyên khoa'],
    ['Doctors', '1-N', 'Appointment', 'Bác sĩ có nhiều lịch hẹn'],
    ['Doctors', '1-N', 'ExaminationRecord', 'Bác sĩ có nhiều phiếu khám'],
    ['PatientProfile', '1-N', 'Appointment', 'Mỗi profile có thể có nhiều lịch hẹn'],
    ['PatientProfile', '1-N', 'ExaminationRecord', 'Mỗi profile có nhiều phiếu khám'],
    ['Appointment', 'N-1', 'Status', 'Mỗi lịch hẹn có 1 trạng thái'],
    ['Appointment', 'N-1', 'Time', 'Mỗi lịch hẹn có 1 khung giờ'],
    ['ExaminationRecord', '1-N', 'CLSIndication', 'Mỗi phiếu khám có nhiều chỉ định CLS'],
    ['CLSIndication', 'N-1', 'CLSService', 'Mỗi chỉ định thuộc 1 dịch vụ CLS'],
]
add_table(doc, rel_headers, rel_rows)

doc.add_page_break()

# ===== SECTION 6: API =====
doc.add_heading('6. API Reference', level=1)
doc.add_paragraph('Hệ thống sử dụng Django views (server-side rendering), không phải REST API. '
                  'Dưới đây là danh sách các endpoint URL chính:')

api_headers = ['Phương thức', 'URL', 'Yêu cầu Auth', 'Request Body', 'Response', 'Ghi chú']
api_rows = [
    ['GET, POST', '/login/', 'Không', 'username, password', 'Redirect dashboard / render login', 'Login'],
    ['GET, POST', '/register/', 'Không', 'user_id, email, password, conf_password,...', 'Redirect / render register', 'Đăng ký BN'],
    ['GET, POST', '/password-reset/', 'Không', 'email', 'Message / render form', 'Quên MK'],
    ['GET, POST', '/reset/<token>/', 'Không', 'password, conf_password', 'Redirect login / render', 'Đặt lại MK'],
    ['GET', '/doctor_dashboard/', 'Doctor', '—', 'Dashboard page', 'Dashboard BS'],
    ['GET, POST', '/view_appointments/', 'Doctor', 'app, status', 'Appointment list', 'QL lịch hẹn'],
    ['GET, POST', '/doctor_exam_create/<id>/', 'Doctor', 'vitals, diagnosis, cls_services', 'Redirect detail', 'Tạo phiếu khám'],
    ['GET, POST', '/doctor_exam_detail/<id>/', 'Doctor', '—', 'Exam detail page', 'Chi tiết PK'],
    ['GET, POST', '/cls_pending/', 'Doctor', '—', 'CLS pending list', 'CLS chờ KQ'],
    ['GET, POST', '/cls_update/<id>/', 'Doctor', 'result', 'Pending list redirect', 'Nhập KQ CLS'],
    ['GET, POST', '/book_appointment/', 'Patient', '—', 'Booking page', 'Đặt lịch'],
    ['GET, POST', '/confirm_book/<doctor>/', 'Patient', 'profile_id, time, date, summary', 'Confirm page', 'Xác nhận đặt'],
    ['GET', '/receptionist/', 'Receptionist', '—', 'Dashboard', 'Dashboard LT'],
    ['GET, POST', '/receptionist/appointments/', 'Receptionist', 'app_id, action', 'Appointment list', 'QL lịch hẹn'],
    ['GET, POST', '/receptionist/walkin/', 'Receptionist', 'full_name, phone, doctor_id,...', 'Redirect', 'Đăng ký trực tiếp'],
    ['GET', '/administration/', 'Admin', '—', 'Dashboard page', 'Dashboard QT'],
    ['GET, POST', '/administration/users/', 'Admin', '—', 'User list', 'DS tài khoản'],
    ['GET, POST', '/administration/categories/', 'Admin', 'name, type, description', 'Categories', 'QL danh mục'],
    ['GET', '/administration/reports/', 'Admin', '—', 'Reports page', 'Báo cáo'],
    ['GET', '/admin/', 'Staff', '—', 'Django admin', 'Admin built-in'],
]
add_table(doc, api_headers, api_rows)

doc.add_page_break()

# ===== SECTION 7: CONFIG =====
doc.add_heading('7. Cấu hình & Môi trường', level=1)

doc.add_heading('7.1. Biến môi trường (.env)', level=2)
env_headers = ['Biến', 'Mục đích', 'Giá trị mẫu', 'Bắt buộc']
env_rows = [
    ['SECRET_KEY', 'Khóa bí mật Django (sinh bằng django.core.management.utils.get_random_secret_key)', 'django-insecure-abc...', 'Bắt buộc'],
    ['DEBUG', 'Chế độ debug (True cho dev, False cho production)', 'True', 'Tùy chọn'],
    ['ALLOWED_HOSTS', 'Danh sách host được phép (phân cách bằng dấu phẩy)', 'localhost,127.0.0.1', 'Tùy chọn'],
    ['DB_NAME', 'Tên database MySQL', 'pbmed', 'Bắt buộc'],
    ['DB_USER', 'User MySQL', 'root', 'Bắt buộc'],
    ['DB_PASSWORD', 'Password MySQL', '', 'Tùy chọn'],
    ['DB_HOST', 'Host MySQL', 'localhost', 'Tùy chọn'],
    ['DB_PORT', 'Port MySQL', '3306', 'Tùy chọn'],
    ['EMAIL_HOST', 'SMTP server', 'smtp.gmail.com', 'Không bắt buộc (fallback)'],
    ['EMAIL_PORT', 'SMTP port', '587', 'Không bắt buộc'],
    ['EMAIL_USE_TLS', 'Dùng TLS?', 'True', 'Không bắt buộc'],
    ['EMAIL_HOST_USER', 'Email gửi', 'your-email@gmail.com', 'Không bắt buộc'],
    ['EMAIL_HOST_PASSWORD', 'App password', '', 'Không bắt buộc'],
]
add_table(doc, env_headers, env_rows)

doc.add_heading('7.2. Cấu hình Django (settings.py)', level=2)
cfg_headers = ['Cấu hình', 'Giá trị', 'Ảnh hưởng']
cfg_rows = [
    ['AUTH_USER_MODEL', 'users.Users', 'Custom User model thay vì auth.User mặc định'],
    ['INSTALLED_APPS', '5 apps + 7 Django built-in', 'Các module được kích hoạt'],
    ['DATABASES', 'MySQL + utf8mb4', 'Kết nối CSDL với charset tiếng Việt'],
    ['TEMPLATES.DIRS', 'BASE_DIR/templates', 'Thư mục chứa template gốc'],
    ['STATICFILES_DIRS', 'BASE_DIR/static', 'Thư mục chứa file tĩnh'],
    ['STATIC_ROOT', 'BASE_DIR/assets', 'Nơi collectstatic xuất file'],
    ['MEDIA_ROOT', 'BASE_DIR/media', 'Nơi lưu file upload'],
    ['EMAIL_BACKEND', 'smtp.EmailBackend', 'Gửi email SMTP'],
]
add_table(doc, cfg_headers, cfg_rows)

doc.add_page_break()

# ===== SECTION 8: LIMITATIONS =====
doc.add_heading('8. Hạn chế & Nợ kỹ thuật', level=1)

doc.add_heading('8.1. Bảo mật', level=2)
doc.add_paragraph('• Mật khẩu mặc định "123456" cho tất cả tài khoản seed — cần đổi mật khẩu khi triển khai thật. '
                  'Đây là lựa chọn có chủ đích để dễ test, nhưng không an toàn cho production.', style='List Bullet')
doc.add_paragraph('• DEBUG=True mặc định — lộ thông tin lỗi chi tiết. Cần set DEBUG=False + ALLOWED_HOSTS khi deploy.', style='List Bullet')
doc.add_paragraph('• Email server chưa được cấu hình — forgot password fallback về hiển thị link trên màn hình.', style='List Bullet')
doc.add_paragraph('• Không có rate limiting cho login/forgot — dễ bị brute-force attack.', style='List Bullet')

doc.add_heading('8.2. Kiến trúc', level=2)
doc.add_paragraph('• Address model (users/models.py) và id_address field, profile_avatar field đã được xoá trong migration 0002 — '
                  'không còn di tích trong code.', style='List Bullet')
doc.add_paragraph('• Một số view thiếu select_related/prefetch_related — gây N+1 query trên production.', style='List Bullet')

doc.add_heading('8.3. Tính năng', level=2)
doc.add_paragraph('• Không có tìm kiếm nâng cao — chỉ hỗ trợ filter cơ bản bằng icontains.', style='List Bullet')
doc.add_paragraph('• Không có phân trang — danh sách user/appointment dễ bị quá tải khi có nhiều dữ liệu.', style='List Bullet')
doc.add_paragraph('• Không có lịch sử thay đổi (audit log) — không track ai đã sửa gì.', style='List Bullet')
doc.add_paragraph('• Không có thông báo real-time — bệnh nhân không nhận được thông báo khi lịch được xác nhận.', style='List Bullet')
doc.add_paragraph('• Chưa có module thuốc/dược — phiếu khám chỉ có chẩn đoán và CLS, không có kê đơn.', style='List Bullet')

doc.add_heading('8.4. Giao diện', level=2)
doc.add_paragraph('• Template chưa responsive hoàn toàn trên mobile.', style='List Bullet')
doc.add_paragraph('• Thiếu validation phía client (JavaScript) — dễ gây trải nghiệm kém khi nhập sai.', style='List Bullet')
doc.add_paragraph('• Một số form không có thông báo lỗi ở UI, chỉ có Django messages ở đầu trang.', style='List Bullet')

doc.add_heading('8.5. Lý do chấp nhận trade-off', level=2)
doc.add_paragraph(
    'Các hạn chế trên được chấp nhận vì đây là đồ án học thuật (Đồ án I) với mục tiêu chính là chứng minh '
    'kiến trúc 3 tầng và quy trình nghiệp vụ. Code được ưu tiên tính rõ ràng, dễ hiểu và đầy đủ chức năng '
    'core, thay vì tối ưu performance hay bảo mật production. Việc sử dụng mật khẩu mặc định và thiếu audit log '
    'là những đánh đổi có chủ đích để giảm độ phức tạp khi demo.'
)

# Save
doc.save('PBMed_TaiLieuHeThong.docx')
print(f'Done: PBMed_TaiLieuHeThong.docx')
